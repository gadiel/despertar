#!/usr/bin/env python3
"""Generate El Libro del Despertar DOCX with original styling."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

# === COLORS ===
C_TITLE = RGBColor(0x1A, 0x1A, 0x1A)
C_ORNAMENT = RGBColor(0xAA, 0xAA, 0xAA)
C_SUBTITLE = RGBColor(0x55, 0x55, 0x55)
C_BOOK_TITLE = RGBColor(0x2F, 0x2F, 0x2F)
C_VERSE_NUM = RGBColor(0x8B, 0x45, 0x13)  # Saddle brown
C_BODY = RGBColor(0x1A, 0x1A, 0x1A)
C_SEPARATOR = RGBColor(0x99, 0x99, 0x99)
C_EPIGRAPH = RGBColor(0x55, 0x55, 0x55)
C_POEM = RGBColor(0x55, 0x55, 0x55)

# === FONTS ===
F_TITLE = "Garamond"
F_BODY = "Georgia"

# === SIZES ===
S_MAIN_TITLE = Pt(22)
S_ORNAMENT = Pt(10)
S_SUBTITLE_COVER = Pt(13)
S_EPIGRAPH_COVER = Pt(11)
S_BOOK_TITLE = Pt(13)
S_BOOK_SUBTITLE = Pt(11)
S_EPIGRAPH = Pt(11)
S_BODY = Pt(11.5)
S_SEPARATOR = Pt(12)
S_SECTION_HDR = Pt(12)
S_POEM = Pt(11)
S_DIALOGUE = Pt(11.5)


def set_spacing(paragraph, before=0, after=0, line=None):
    pPr = paragraph._p.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = pPr.makeelement(qn('w:spacing'), {})
        pPr.append(sp)
    if before:
        sp.set(qn('w:before'), str(before))
    if after:
        sp.set(qn('w:after'), str(after))
    if line:
        sp.set(qn('w:line'), str(line))
        sp.set(qn('w:lineRule'), 'auto')


def run(p, text, font=F_BODY, size=S_BODY, bold=False, italic=False, color=C_BODY):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = size
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    # Set east-asia and complex-script fonts too
    rPr = r._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'), font)
    return r


def add_inline(p, text, font=F_BODY, size=S_BODY, color=C_BODY, base_bold=False, base_italic=False):
    """Parse markdown inline formatting (*italic*, **bold**, ***both***) and add runs."""
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    for m in re.finditer(pattern, text):
        if m.group(2):  # ***bold+italic***
            run(p, m.group(2), font, size, bold=True, italic=True, color=color)
        elif m.group(3):  # **bold**
            run(p, m.group(3), font, size, bold=True, italic=base_italic, color=color)
        elif m.group(4):  # *italic*
            run(p, m.group(4), font, size, bold=base_bold, italic=True, color=color)
        elif m.group(5):  # plain
            run(p, m.group(5), font, size, bold=base_bold, italic=base_italic, color=color)


def add_title_page(doc):
    # Vertical space to push content down
    for _ in range(6):
        p = doc.add_paragraph()
        set_spacing(p, after=0)

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=200)
    run(p, "EL LIBRO DEL DESPERTAR", F_TITLE, S_MAIN_TITLE, bold=True, color=C_TITLE)

    # Ornamental line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=100)
    run(p, "───────────────", F_BODY, S_ORNAMENT, color=C_ORNAMENT)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=600)
    run(p, "Escrituras para la persona de hoy", F_BODY, S_SUBTITLE_COVER, italic=True, color=C_SUBTITLE)

    # Epigraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=200, after=300)
    run(p, "Lo que antes fue zarza ardiente, hoy es la pregunta\nque no te deja dormir a las tres de la mañana.",
        F_BODY, S_EPIGRAPH_COVER, italic=True, color=C_SUBTITLE)


def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=300, after=300)
    run(p, "⸻", F_BODY, S_SEPARATOR, color=C_SEPARATOR)


def add_book_title(doc, title, page_break=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if page_break:
        p.paragraph_format.page_break_before = True
    set_spacing(p, before=480, after=280)
    run(p, title.upper(), F_TITLE, S_BOOK_TITLE, bold=True, color=C_BOOK_TITLE)


def add_book_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=200)
    clean = text.strip('()').strip()
    run(p, f"({clean})", F_BODY, S_BOOK_SUBTITLE, italic=True, color=C_SUBTITLE)


def add_epigraph_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=100, after=100)
    run(p, text, F_BODY, S_EPIGRAPH, italic=True, color=C_EPIGRAPH)


def add_verse(doc, num, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, after=180, line=320)
    # Verse number in saddle brown
    run(p, f"{num}. ", F_BODY, S_BODY, bold=True, italic=True, color=C_VERSE_NUM)
    # Body text with inline formatting
    add_inline(p, text, F_BODY, S_BODY, C_BODY)


def add_section_header(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=360, after=240)
    run(p, title, F_BODY, S_SECTION_HDR, italic=True, color=C_VERSE_NUM)


def add_narrative(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, after=180, line=320)
    add_inline(p, text, F_BODY, S_BODY, C_BODY)


def add_poem_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.5)
    if text.strip():
        set_spacing(p, after=40, line=280)
        add_inline(p, text, F_BODY, S_POEM, C_POEM, base_italic=True)
    else:
        set_spacing(p, after=80, line=140)


def add_dialogue_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, after=120, line=320)
    p.paragraph_format.left_indent = Inches(0.3)
    add_inline(p, text, F_BODY, S_DIALOGUE, C_BODY)


# === MARKDOWN PARSER ===

def parse_file(filepath):
    """Parse a markdown file into structured elements."""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = [l.rstrip('\n') for l in f.readlines()]

    elements = []
    section = None  # 'parabola', 'salmo', 'decima', 'dialogo', None
    seen_first_sep = False

    for line in lines:
        stripped = line.strip()

        # Empty line
        if not stripped:
            if section in ('salmo', 'decima'):
                elements.append(('poem_blank', ''))
            continue

        # Book title
        m = re.match(r'^# (.+)$', stripped)
        if m:
            elements.append(('title', m.group(1)))
            continue

        # Book subtitle
        m = re.match(r'^## (.+)$', stripped)
        if m:
            elements.append(('subtitle', m.group(1)))
            continue

        # Separator
        if stripped == '---':
            if not seen_first_sep:
                seen_first_sep = True
            else:
                elements.append(('separator', ''))
                section = None
            continue

        # Epigraph (lines before first ---)
        if not seen_first_sep:
            # Strip markdown italic markers
            clean = stripped
            if clean.startswith('*') and clean.endswith('*') and not clean.startswith('**'):
                clean = clean[1:-1]
            elements.append(('epigraph', clean))
            continue

        # Section header (### ...)
        m = re.match(r'^### (.+)$', stripped)
        if m:
            header = m.group(1)
            lo = header.lower()
            if 'parábola' in lo or 'parabola' in lo:
                section = 'parabola'
            elif 'salmo' in lo or 'canto' in lo:
                section = 'salmo'
            elif 'décima' in lo or 'decima' in lo:
                section = 'decima'
            elif 'diálogo' in lo or 'dialogo' in lo or 'disputa' in lo:
                section = 'dialogo'
            else:
                section = 'parabola'  # default subsection
            elements.append(('section_header', header))
            continue

        # Numbered verse: **N.** text
        m = re.match(r'^\*\*(\d+)\.\*\*\s*(.*)$', stripped)
        if m:
            section = None
            elements.append(('verse', (m.group(1), m.group(2))))
            continue

        # Context-dependent content
        if section in ('salmo', 'decima'):
            elements.append(('poem_line', stripped))
        elif section == 'dialogo':
            elements.append(('dialogue', stripped))
        elif section == 'parabola':
            elements.append(('narrative', stripped))
        elif stripped.startswith('—'):
            elements.append(('dialogue', stripped))
        else:
            elements.append(('narrative', stripped))

    return elements


def build_docx(files, output):
    doc = Document()

    # Default style
    style = doc.styles['Normal']
    style.font.name = F_BODY
    style.font.size = S_BODY
    style.font.color.rgb = C_BODY

    # Page setup
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1.25)
    sec.bottom_margin = Inches(1.25)
    sec.left_margin = Inches(1.25)
    sec.right_margin = Inches(1.25)

    # Title page
    add_title_page(doc)

    for filepath in files:
        elems = parse_file(filepath)
        first_title = True

        for etype, content in elems:
            if etype == 'title':
                add_book_title(doc, content, page_break=True)
                first_title = False
            elif etype == 'subtitle':
                add_book_subtitle(doc, content)
            elif etype == 'epigraph':
                add_epigraph_line(doc, content)
            elif etype == 'separator':
                add_separator(doc)
            elif etype == 'verse':
                num, text = content
                add_verse(doc, num, text)
            elif etype == 'section_header':
                add_section_header(doc, content)
            elif etype == 'narrative':
                add_narrative(doc, content)
            elif etype == 'poem_line':
                add_poem_line(doc, content)
            elif etype == 'poem_blank':
                add_poem_line(doc, '')
            elif etype == 'dialogue':
                add_dialogue_line(doc, content)

    doc.save(output)
    print(f"Generated: {output} ({Path(output).stat().st_size // 1024} KB)")


# === FILE ORDER ===
BASE = Path('/Users/gadiel/despertar/libros')
FILES = [
    '00-prologo-la-zarza-que-nadie-ve.md',
    '01-del-ruido-y-el-silencio.md',
    '02-libro-segundo-de-la-soledad-y-el-projimo.md',
    '03-del-trabajo-y-el-sentido.md',
    '04-de-la-muerte-y-la-finitud.md',
    '05-de-la-fe-sin-nombre.md',
    '06-del-padre-y-la-hija.md',
    '07-del-pan-y-la-tierra.md',
    '08-del-dinero-y-la-deuda.md',
    '09-de-la-ansiedad-y-la-sanacion.md',
    '10-del-amor-y-el-cuerpo.md',
    '11-de-la-fiesta-y-el-asombro.md',
    '12-del-camino-y-la-busqueda.md',
    '13-de-la-justicia-y-el-extranjero.md',
    '14-de-las-cadenas-y-la-libertad.md',
    '15-del-espejo-y-la-maquina.md',
    '16-de-la-mesa-compartida.md',
    '17-epilogo-el-desierto-es-ahora.md',
]

if __name__ == '__main__':
    paths = [str(BASE / f) for f in FILES]
    # Verify all files exist
    for p in paths:
        if not Path(p).exists():
            print(f"ERROR: File not found: {p}")
            exit(1)
    build_docx(paths, '/Users/gadiel/despertar/El_Libro_del_Despertar.docx')
